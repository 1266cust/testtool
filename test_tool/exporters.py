from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import List

import csv
from openpyxl import Workbook
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from .core.models import TestCase


THUMBNAIL_PNG = (
    b"\x89PNG\r\n\x1a\n"
    b"\x00\x00\x00\rIHDR"
    b"\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89"
    b"\x00\x00\x00\x0cIDAT"
    b"\x08\xd7c\xf8\xff\xff?\x00\x05\xfe\x02\xfe\xa7\x9d\xa4\x9a"
    b"\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _with_timestamp_suffix(path: Path) -> Path:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    return path.with_name(f"{path.stem}_{ts}{path.suffix}")


HEADERS = [
    "用例ID",
    "用例模块",
    "用例名称",
    "验收目的",
    "预置条件",
    "操作步骤",
    "预期结果",
    "用例类型",
    "优先级",
]


def export_cases_to_excel(cases: List[TestCase], path: Path) -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "TestCases"

    ws.append(HEADERS)

    for c in cases:
        ws.append(
            [
                c.case_id,
                c.module,
                c.name,
                c.acceptance_purpose,
                c.preconditions,
                c.steps,
                c.expected_result,
                c.case_type,
                c.priority,
            ]
        )

    for col in range(1, len(HEADERS) + 1):
        ws.column_dimensions[chr(64 + col)].width = 20

    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        wb.save(str(path))
        return path
    except PermissionError:
        alt = _with_timestamp_suffix(path)
        wb.save(str(alt))
        return alt


def export_cases_to_word(cases: List[TestCase], path: Path, title: str = "测试用例") -> Path:
    """导出测试用例到Word文档 - 一个用例一个竖向表格"""
    from docx.oxml.ns import qn

    def set_chinese_font(run, font_name: str = "微软雅黑", size: int = 10):
        """设置中文字体"""
        run.font.name = font_name
        run.font.size = Pt(size)
        # 设置东亚字体
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), font_name)

    def add_field_row(table, row_idx: int, field_name: str, value: str):
        """添加一行：字段名 | 值"""
        row = table.rows[row_idx]
        # 字段名单元格
        cell0 = row.cells[0]
        cell0.text = ""
        p0 = cell0.paragraphs[0]
        run0 = p0.add_run(field_name)
        run0.bold = True
        set_chinese_font(run0, "微软雅黑", 10)

        # 值单元格
        cell1 = row.cells[1]
        cell1.text = ""
        p1 = cell1.paragraphs[0]
        run1 = p1.add_run(str(value) if value else "")
        set_chinese_font(run1, "微软雅黑", 10)

    doc = Document()

    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 文档标题
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        set_chinese_font(run, "微软雅黑", 18)

    # 生成信息
    info_para = doc.add_paragraph()
    run_time = info_para.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    set_chinese_font(run_time, "微软雅黑", 10)
    run_count = info_para.add_run(f"    用例数量：{len(cases)} 条")
    set_chinese_font(run_count, "微软雅黑", 10)

    doc.add_paragraph()

    # 字段定义：(字段名, 字段显示名称)
    fields = [
        ("case_id", "用例ID"),
        ("module", "用例模块"),
        ("name", "用例名称"),
        ("acceptance_purpose", "验收目的"),
        ("preconditions", "预置条件"),
        ("steps", "操作步骤"),
        ("expected_result", "预期结果"),
        ("case_type", "用例类型"),
        ("priority", "优先级"),
    ]

    # 每个用例一个竖向表格
    for idx, c in enumerate(cases):
        # 用例编号标题
        case_title = doc.add_paragraph()
        run_title = case_title.add_run(f"【用例 {idx + 1}】")
        run_title.bold = True
        set_chinese_font(run_title, "微软雅黑", 11)

        # 创建 2列 x N行 的表格
        table = doc.add_table(rows=len(fields), cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # 设置列宽
        for row in table.rows:
            row.cells[0].width = Inches(1.2)
            row.cells[1].width = Inches(5.3)

        # 填充数据
        for i, (field_name, display_name) in enumerate(fields):
            value = getattr(c, field_name, "")
            # 列表类型转字符串
            if isinstance(value, list):
                value = "\n".join(f"{j+1}. {v}" for j, v in enumerate(value)) if value else ""
            add_field_row(table, i, display_name, value)

        # 用例之间添加空行
        doc.add_paragraph()

    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(path))
        return path
    except PermissionError:
        alt = _with_timestamp_suffix(path)
        doc.save(str(alt))
        return alt


def export_cases_to_csv(cases: List[TestCase], path: Path) -> Path:
    rows = [
        [
            c.case_id,
            c.module,
            c.name,
            c.acceptance_purpose,
            c.preconditions,
            c.steps,
            c.expected_result,
            c.case_type,
            c.priority,
        ]
        for c in cases
    ]
    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        with path.open("w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(HEADERS)
            writer.writerows(rows)
        return path
    except PermissionError:
        alt = _with_timestamp_suffix(path)
        with alt.open("w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(HEADERS)
            writer.writerows(rows)
        return alt


def export_cases_to_xmind(cases: List[TestCase], path: Path, root_title: str) -> Path:
    """
    生成 XMind Zen/2026 兼容的 `.xmind` 文件（zip + content.json）。
    """
    import json
    import uuid
    import zipfile

    def new_id() -> str:
        return uuid.uuid4().hex

    def topic(title: str, children: List[dict] | None = None) -> dict:
        t: dict = {"id": new_id(), "class": "topic", "title": title}
        if children:
            t["children"] = {"attached": children}
        return t

    modules: dict = {}
    for c in cases:
        modules.setdefault(c.module, {}).setdefault(c.case_type, []).append(c)

    module_topics: List[dict] = []
    for module_name, sub_map in modules.items():
        sub_topics: List[dict] = []
        for sub_name, sub_cases in sub_map.items():
            case_topics = [
                topic(f"{case.case_id} - {case.name}")
                for case in sub_cases
            ]
            sub_topics.append(topic(sub_name, case_topics))
        module_topics.append(topic(module_name, sub_topics))

    now_ms = int(datetime.now(tz=timezone.utc).timestamp() * 1000)
    sheet_id = new_id()
    root_topic = topic(root_title, module_topics)
    root_topic["structureClass"] = "org.xmind.ui.map.unbalanced"

    content = [
        {
            "id": sheet_id,
            "class": "sheet",
            "title": root_title,
            "rootTopic": root_topic,
        }
    ]

    metadata = {}

    manifest = {
        "file-entries": {
            "content.json": {},
            "metadata.json": {},
            "Thumbnails/thumbnail.png": {},
        }
    }

    path.parent.mkdir(parents=True, exist_ok=True)

    def write_zip(target: Path) -> None:
        with zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("resources/", b"")
            zf.writestr("Thumbnails/", b"")
            zf.writestr("content.json", json.dumps(content, ensure_ascii=False))
            zf.writestr("metadata.json", json.dumps(metadata, ensure_ascii=False))
            zf.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False))
            zf.writestr("Thumbnails/thumbnail.png", THUMBNAIL_PNG)

            import xml.etree.ElementTree as ET

            NS_CONTENT = "urn:xmind:xmap:xmlns:content:2.0"
            NS_FO = "http://www.w3.org/1999/XSL/Format"
            NS_SVG = "http://www.w3.org/2000/svg"
            NS_XHTML = "http://www.w3.org/1999/xhtml"
            NS_XLINK = "http://www.w3.org/1999/xlink"

            ET.register_namespace("", NS_CONTENT)
            ET.register_namespace("fo", NS_FO)
            ET.register_namespace("svg", NS_SVG)
            ET.register_namespace("xhtml", NS_XHTML)
            ET.register_namespace("xlink", NS_XLINK)

            xmap = ET.Element(
                f"{{{NS_CONTENT}}}xmap-content",
                attrib={
                    "timestamp": str(now_ms),
                    "version": "2.0",
                },
            )

            sheet_el = ET.SubElement(
                xmap,
                f"{{{NS_CONTENT}}}sheet",
                attrib={"id": sheet_id, "timestamp": str(now_ms)},
            )

            def add_topic(parent, t: dict) -> None:
                topic_el = ET.SubElement(parent, f"{{{NS_CONTENT}}}topic", attrib={"id": t["id"]})
                title_el = ET.SubElement(topic_el, f"{{{NS_CONTENT}}}title")
                title_el.text = t.get("title", "")

                children = t.get("children", {}).get("attached", [])
                if children:
                    children_el = ET.SubElement(topic_el, f"{{{NS_CONTENT}}}children")
                    topics_el = ET.SubElement(children_el, f"{{{NS_CONTENT}}}topics", attrib={"type": "attached"})
                    for child in children:
                        add_topic(topics_el, child)

            add_topic(sheet_el, content[0]["rootTopic"])

            content_xml_bytes = ET.tostring(xmap, encoding="utf-8", xml_declaration=True)
            zf.writestr("content.xml", content_xml_bytes)

    try:
        write_zip(path)
        return path
    except PermissionError:
        alt = _with_timestamp_suffix(path)
        write_zip(alt)
        return alt


def export_cases_to_freemind(cases: List[TestCase], path: Path, root_title: str) -> Path:
    """
    生成一个简单的 Freemind `.mm` 文件，可导入 XMind。
    """
    import xml.etree.ElementTree as ET

    path.parent.mkdir(parents=True, exist_ok=True)

    TEXT = "TEXT"

    map_el = ET.Element("map", version="1.0.1")
    root_node = ET.SubElement(map_el, "node", TEXT=root_title)

    modules = {}
    for c in cases:
        modules.setdefault(c.module, {}).setdefault(c.case_type, []).append(c)

    for module_name, sub_map in modules.items():
        module_node = ET.SubElement(root_node, "node", TEXT=module_name)
        for sub_name, sub_cases in sub_map.items():
            sub_node = ET.SubElement(module_node, "node", TEXT=sub_name)
            for case in sub_cases:
                case_text = f"{case.case_id} - {case.name}"
                ET.SubElement(sub_node, "node", TEXT=case_text)

    tree = ET.ElementTree(map_el)
    try:
        tree.write(str(path), encoding="utf-8", xml_declaration=True)
        return path
    except PermissionError:
        alt = _with_timestamp_suffix(path)
        tree.write(str(alt), encoding="utf-8", xml_declaration=True)
        return alt